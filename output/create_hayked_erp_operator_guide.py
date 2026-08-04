from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = r"D:\antigravity project\Hayked coffee ERP\output\Hayked Coffee Warehouse ERP - Real World Flow Guide.docx"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def borders(table, color="D9E2EA"):
    tbl_pr = table._tbl.tblPr
    tbl_borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tbl_borders.append(el)
    tbl_pr.append(tbl_borders)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table(doc, headers, rows, widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    borders(tbl)
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        shade(cell, "E8EEF5")
        set_cell_text(cell, h, bold=True, color="0B2545")
        if widths:
            cell.width = Inches(widths[i])
    for row in rows:
        cells = tbl.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph()
    return tbl


def h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(13, 89, 112)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def callout(doc, title, body):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    borders(tbl, "B9DCE6")
    cell = tbl.rows[0].cells[0]
    shade(cell, "F3FBFD")
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor(0, 112, 132)
    p.add_run("\n" + body)
    doc.add_paragraph()


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(0.8)
section.right_margin = Inches(0.8)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(10.5)
styles["Normal"].paragraph_format.space_after = Pt(6)
styles["Normal"].paragraph_format.line_spacing = 1.15
for name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 12)]:
    styles[name].font.name = "Calibri"
    styles[name].font.size = Pt(size)
    styles[name].font.bold = True
    styles[name].paragraph_format.space_before = Pt(10)
    styles[name].paragraph_format.space_after = Pt(5)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Hayked Coffee Warehouse ERP")
r.bold = True
r.font.size = Pt(24)
r.font.color.rgb = RGBColor(0, 56, 76)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Real-world training story, feature guide, and sample data for first pilot testing")
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(82, 99, 120)

callout(
    doc,
    "How to use this guide",
    "Follow the story of Alem Export PLC from client onboarding to GRN, stock, processing, dispatch, finance, documents, approvals, audit, and administration. Use the sample values exactly while testing the ERP.",
)

h(doc, "1. What the software is", 1)
doc.add_paragraph(
    "Hayked Coffee Warehouse ERP is an operations system for a coffee processing and warehousing company. "
    "It controls client-owned coffee from the first truck arrival until the coffee is processed, dispatched, billed, documented, and audited."
)
doc.add_paragraph(
    "The main idea is simple: every important warehouse action should create a record, every stock movement should balance in kilograms, and every sensitive action should require the right approval."
)

h(doc, "2. The real-world story: Alem Export PLC", 1)
doc.add_paragraph(
    "Alem Export PLC is a coffee exporter. Alem does not own a full warehouse and processing facility, so she brings coffee to Hayked. "
    "Hayked stores the coffee, protects ownership, processes it for export preparation, controls rejects and byproducts, releases coffee only after finance and document checks, and gives Alem traceable records."
)
table(
    doc,
    ["Person / company", "Role in the story", "Why it matters"],
    [
        ["Alem Export PLC", "Client/customer", "Owns the coffee and pays Hayked for warehouse and processing services."],
        ["Bekele Alemu", "Authorized representative", "Allowed to deliver and collect coffee on behalf of Alem."],
        ["Meron Tadesse", "Warehouse manager", "Approves sensitive warehouse movements and supervises operations."],
        ["Abebe Tesfaye", "Warehouse officer", "Creates GRNs and day-to-day warehouse records."],
        ["Hana Bekele", "Processing supervisor", "Approves processing requests and completion records."],
        ["Saba Getachew", "Finance officer", "Issues invoices and records payments before dispatch release."],
    ],
    [1.35, 1.65, 3.5],
)

h(doc, "3. First sample master data", 1)
doc.add_paragraph("Enter this data first. Without client, agreement, and representative records, the warehouse flow should not be trusted.")
table(
    doc,
    ["Feature", "Sample data to enter", "Business meaning"],
    [
        ["Client", "Code: ALEM; Legal name: Alem Export PLC; TIN: 0099887766; Phone: 0911000000; Email: operations@alemexport.example", "Defines who owns the coffee."],
        ["Agreement", "Agreement no: AGR-ALEM-2026-001; Effective from: 2026-08-02; Status: ACTIVE; Tariff version: HAYKED-2026-A; Default bag kg: 60", "Defines valid service terms."],
        ["Representative", "Name: Bekele Alemu; ID: REP-ALEM-001; Phone: 0911222333; Valid from: 2026-08-02; Active: yes", "Defines who can represent the client."],
    ],
    [1.2, 3.55, 1.75],
)

h(doc, "4. Feature-by-feature guide", 1)
features = [
    ("Dashboard", "Gives the manager a quick view of total coffee, arrivals, processed coffee, rejects, queued work, dispatches, and operational alerts.", "Before work starts, Meron checks whether today has pending approvals, low bag stock, or processing exceptions."),
    ("Clients", "Stores customers, agreements, and authorized representatives.", "Alem must exist here before her coffee can be received properly."),
    ("Warehouse / GRN", "Records coffee arrival and creates official warehouse receipts.", "Alem sends 100 bags of washed coffee by truck. Hayked creates a GRN and later posts it into stock."),
    ("Stock / Lots", "Tracks each coffee lot by owner, location, quantity, status, and history.", "After GRN posting, LOT-ALEM-0001 becomes the traceable stock record."),
    ("Processing Requests", "Digitizes the paper Export Coffee Processing Order Requesting Notes form.", "Alem requests 80 bags to be prepared for export. The request must be approved before queueing."),
    ("Processing Queue", "Controls which approved lots are waiting for processing.", "Hana moves Alem's approved request into the queue."),
    ("Processing Orders", "Tracks the actual processing job, input kg, output kg, rejects, byproduct, and loss.", "The system checks that 4,800 kg input equals all output categories."),
    ("Dispatch", "Releases coffee only when approval, finance, documents, weighbridge, and hold checks are satisfied.", "Alem cannot collect coffee until the dispatch gates are complete."),
    ("Finance", "Issues invoices and records payments or credit approvals.", "Saba confirms payment/credit before dispatch is posted."),
    ("Documents", "Stores scans and files linked to GRNs, requests, invoices, dispatches, and approvals.", "The scanned paper request note and dispatch documents stay attached to records."),
    ("Approvals", "Prevents the same person from requesting and approving sensitive actions.", "A requester cannot approve their own processing request."),
    ("Audit", "Keeps traceability of important actions.", "When someone asks what happened to Alem's coffee, Hayked can show a trail."),
    ("Administration", "Controls users, roles, warehouses, and system setup.", "Start broad as admin during testing, then later restrict users by real duties."),
]
table(doc, ["Feature", "What it is used for", "Alem story example"], features, [1.35, 2.6, 2.55])

h(doc, "5. End-to-end test flow with exact sample entries", 1)
h(doc, "Step 1: Add Alem as a client", 2)
for item in [
    "Open Clients or Administration > Client setup.",
    "Create client ALEM / Alem Export PLC.",
    "Add active agreement AGR-ALEM-2026-001.",
    "Add representative Bekele Alemu.",
]:
    numbered(doc, item)
callout(doc, "Expected result", "The ERP now knows who owns the coffee and who is allowed to act for that client.")

h(doc, "Step 2: Receive coffee with GRN", 2)
table(
    doc,
    ["Field", "Value"],
    [
        ["GRN number", "GRN-2026-0001"],
        ["Client", "Alem Export PLC"],
        ["Representative", "Bekele Alemu"],
        ["Coffee type", "Washed"],
        ["Bags", "100"],
        ["Net kg", "6,000"],
        ["Vehicle plate", "AA-12345"],
        ["Driver", "Dawit Kebede"],
        ["Section", "A-01 Arrival"],
        ["Moisture", "10.8%"],
        ["Origin", "Sidama"],
        ["Grade", "Grade 2"],
    ],
    [2.0, 4.2],
)
for item in ["Save as draft.", "Submit the GRN.", "Approve the GRN using a different approver.", "Post it to stock with lot number LOT-ALEM-0001."]:
    numbered(doc, item)
callout(doc, "Expected result", "A stock lot is created: LOT-ALEM-0001, 100 bags, 6,000 kg, owned by Alem Export PLC.")

h(doc, "Step 3: Create processing request from paper form", 2)
table(
    doc,
    ["Field", "Value"],
    [
        ["Request note number", "PRN-00239"],
        ["Request date", "2026-08-02"],
        ["Client/customer", "Alem Export PLC"],
        ["Warehouse receipt or lot", "LOT-ALEM-0001"],
        ["Coffee type", "Washed"],
        ["Requested preparation type", "Export preparation / Grade 2 screen prep"],
        ["Grade", "Grade 2"],
        ["Requested bags", "80"],
        ["Requested kg", "4,800"],
        ["Certification", "Organic, RFA"],
        ["Representative/requester", "Meron Tadesse"],
        ["Checker", "Abebe Tesfaye"],
        ["Approver", "Hana Bekele"],
        ["Scanned document attached", "Yes"],
        ["Notes", "Prepare for export shipment after finance clearance."],
    ],
    [2.15, 4.05],
)
for item in ["Save as draft.", "Submit request.", "Approve request as a different approver.", "Move approved request to Processing Queue."]:
    numbered(doc, item)
callout(doc, "Expected result", "The request cannot enter queue until status is APPROVED. This protects the warehouse from processing unofficial requests.")

h(doc, "Step 4: Process the coffee", 2)
table(
    doc,
    ["Processing value", "Sample amount", "Meaning"],
    [
        ["Input", "4,800 kg", "Coffee issued from Alem's lot into processing."],
        ["Accepted client coffee", "3,700 kg", "Good processed coffee still owned by Alem."],
        ["Client reject", "100 kg", "Reject that remains client-owned."],
        ["Hayked byproduct", "800 kg", "Byproduct kept by Hayked if agreement allows."],
        ["Process loss", "200 kg", "Normal processing loss."],
        ["Total output", "4,800 kg", "Must equal input."],
    ],
    [1.9, 1.45, 3.15],
)
callout(doc, "Expected result", "The system accepts the completion because the weights reconcile: 3,700 + 100 + 800 + 200 = 4,800 kg.")

h(doc, "Step 5: Dispatch processed coffee", 2)
table(
    doc,
    ["Field", "Value"],
    [
        ["Dispatch number", "DSP-2026-0001"],
        ["Client", "Alem Export PLC"],
        ["Lot", "LOT-ALEM-0001"],
        ["Representative", "Bekele Alemu"],
        ["Quantity", "3,000 kg"],
        ["Bags", "50"],
        ["Invoices paid", "Yes"],
        ["Documents ready", "Yes"],
        ["Weighbridge ready", "Yes"],
        ["Legal/quality hold", "No"],
    ],
    [2.0, 4.2],
)
for item in ["Create dispatch draft.", "Confirm finance and document gates.", "Approve dispatch with an independent approver.", "Post dispatch."]:
    numbered(doc, item)
callout(doc, "Expected result", "Stock is reduced by 3,000 kg and the dispatch record becomes part of the lot history.")

h(doc, "Step 6: Finance and payment", 2)
table(
    doc,
    ["Field", "Value"],
    [
        ["Invoice number", "INV-2026-0001"],
        ["Client", "Alem Export PLC"],
        ["Tariff version", "HAYKED-2026-A"],
        ["Subtotal", "25,000 ETB"],
        ["Tax", "3,750 ETB"],
        ["Total", "28,750 ETB"],
        ["Payment reference", "CBE-TEST-001"],
        ["Payment amount", "28,750 ETB"],
    ],
    [2.0, 4.2],
)
callout(doc, "Expected result", "Finance can prove whether Alem is cleared for dispatch. Later, the invoice/payment history supports billing disputes and reports.")

h(doc, "6. What each role should test first", 1)
table(
    doc,
    ["Role", "First tests", "Should not be allowed"],
    [
        ["System admin", "Create users, change roles, view all modules.", "Use real data before access rules are decided."],
        ["Warehouse manager", "Approve GRN, reverse simple posted GRN, approve dispatch.", "Approve their own prepared record."],
        ["Warehouse officer", "Create GRN, create dispatch draft, view stock.", "Approve high-control transactions alone."],
        ["Processing supervisor", "Create/approve processing requests, queue approved requests, complete orders.", "Queue draft or rejected requests."],
        ["Finance officer", "Issue invoices, record payments, decide finance approvals.", "Post warehouse stock movements."],
        ["Auditor/viewer", "Read reports, audit trail, documents.", "Create or update operational records."],
    ],
    [1.35, 2.65, 2.5],
)

h(doc, "7. What good testing looks like", 1)
for item in [
    "Try the happy path once using the exact Alem data above.",
    "Try one invalid action: approve your own request. The system should block it.",
    "Try one invalid processing completion where output does not equal input. The system should block it.",
    "Try dispatch without finance/documents ready. The system should block it.",
    "Check dashboard numbers after each major step.",
    "Check audit/documents after each approval or posted transaction.",
]:
    bullet(doc, item)

h(doc, "8. Pilot readiness notes", 1)
doc.add_paragraph(
    "For now, use this system as a controlled local/online pilot. Do not enter real client-sensitive data until user roles, backup practice, document retention, service-role key rotation, and operating procedures are finalized."
)
doc.add_paragraph(
    "The next professional step after a successful Alem test is to create two or three real warehouse scenarios: one washed coffee flow, one unwashed/UG flow, and one dispatch with finance hold. These will reveal whether screens and roles match the real warehouse day."
)

doc.save(OUT)
print(OUT)
